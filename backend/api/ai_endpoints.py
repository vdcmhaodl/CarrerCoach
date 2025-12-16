# api/ai_endpoints.py

from fastapi import APIRouter
from fastapi.responses import JSONResponse, StreamingResponse
import json
import re
import io
from docx import Document
from docx.shared import Pt, RGBColor
# Import từ file config/models mới
from core.models import UserInput, CVAnalysisRequest, CVGenerationRequest, QuestionGenerationRequest
from core.config import GEMINI_MODEL # Import model đã được cấu hình

router = APIRouter()

async def get_gemini_evaluation(user_answer: str):
    """
    Contains the logic to call Gemini.
    Receives text (typed or from STT) and returns JSON.
    """
    prompt_template = f"""Bạn là một huấn luyện viên phỏng vấn chuyên gia có tên CareerCoach. Hãy phân tích đầu vào của người dùng và chỉ trả về một đối tượng JSON hợp lệ (không có markdown, không có văn bản bổ sung).

Nếu đầu vào rõ ràng là một câu trả lời phỏng vấn:
Trả về định dạng chính xác này:
{{
  "type": "evaluation",
  "feedback": "Phân hồi chi tiết về điểm mạnh và yếu, với những gợi ý cụ thể để cải thiện. Hãy trả lời bằng tiếng Việt.",
  "suggested_answer": "Một câu trả lời ví dụ tốt hơn cho câu hỏi này dựa trên câu trả lời của người dùng. Hãy trả lời bằng tiếng Việt. Phải cung cấp câu trả lời gợi ý cho mọi câu hỏi, không bao giờ để trống."
}}

Ngoài ra:
Trả về định dạng này:
{{
  "type": "general_answer",
  "response": "Phản hồi của bạn cho đầu vào của người dùng. Hãy trả lời bằng tiếng Việt."
}}

Đầu vào người dùng:
{user_answer}

CHỈ trả về đối tượng JSON, không có văn bản khác. Luôn bao gồm trường "suggested_answer" khi type là "evaluation"."""
    
    if GEMINI_MODEL is None:
        return JSONResponse(
            status_code=503,
            content={
                "error": "Gemini is not configured. Set env var GOOGLE_API_KEY (or GEMINI_API_KEY) on the server.",
            },
        )

    try:
        response = await GEMINI_MODEL.generate_content_async(prompt_template)
        raw_text = response.text.strip()
        
        print(f"Raw Gemini response: {raw_text[:200]}...")
        
        # Cleanup Markdown and code blocks
        raw_text = re.sub(r'```json\s*', '', raw_text)
        raw_text = re.sub(r'```\s*', '', raw_text)
        raw_text = raw_text.strip()
        
        # Try to find JSON object
        match = re.search(r'(\{[\s\S]*\})', raw_text)
        if match:
            json_str = match.group(1)
            print(f"Extracted JSON: {json_str[:200]}...")
            
            # Fix trailing commas before closing brackets
            json_str = re.sub(r',\s*([\]\}])', r'\1', json_str)
            
            # Try to parse JSON
            ai_data = json.loads(json_str)
            print(f"Successfully parsed: {ai_data.get('type')}")
            return JSONResponse(content=ai_data)
        else:
            print(f"No JSON found in response")
            return JSONResponse(
                content={
                    "type": "general_answer",
                    "response": raw_text[:500] if raw_text else "Please provide a clearer input."
                }
            )
    except json.JSONDecodeError as e:
        print(f"JSON parsing error: {e}")
        return JSONResponse(
            content={
                "type": "evaluation",
                "feedback": "Câu trả lời của bạn cho thấy nỗ lực tốt. Hãy tiếp tục luyện tập và cố gắng trở nên cụ thể hơn với các ví dụ.",
                "suggested_answer": "Cung cấp một câu trả lời có cấu trúc hơn với các ví dụ cụ thể từ kinh nghiệm của bạn."
            }
        )
    except Exception as e:
        print(f"Error calling Gemini: {type(e).__name__}: {e}")
        return JSONResponse(
            content={
                "type": "evaluation",
                "feedback": "Tôi đang xử lý phản hồi của bạn. Vui lòng thử lại sau một lúc.",
                "suggested_answer": "Hãy cân nhắc thêm các chi tiết cụ thể hơn vào câu trả lời của bạn."
            }
        )

@router.post("/gemini")
async def handle_gemini_request(data: UserInput):
    return await get_gemini_evaluation(data.prompt)

@router.post("/analyze-cv")
async def analyze_cv(data: CVAnalysisRequest):
    """
    Analyze CV text and extract: role, skills, experience, pros/cons, learning path
    """
    prompt_template = f"""
    Bạn là chuyên gia hướng dẫn nghề nghiệp và phân tích sơ yếu lý lịch.
 
    Phân tích văn bản sơ yếu lý lịch sau và trích xuất các thông tin chi tiết toàn diện:
 
    Văn bản CV:
    {data.cv_text}
 
    VAI TRÒ MỤC TIÊU (nếu có): {data.role or 'Không xác định'}
    TỔ CHỨC MỤC TIÊU (nếu có): {data.organization or 'Không xác định'}
 
    QUAN TRỌNG: TẤT CẢ nội dung phân tích, mô tả, điểm mạnh, điểm yếu, nhiệm vụ đề xuất PHẢI viết bằng TIẾNG VIỆT, kể cả khi CV gốc bằng tiếng Anh.
 
    Cung cấp phân tích chi tiết theo định dạng JSON với cấu trúc sau:
    {{
      "extracted_role": "Vai trò/vị trí chính dựa trên CV (ví dụ: 'Kỹ sư phần mềm', 'Quản lý tiếp thị')",
      "skills": ["kỹ năng1", "kỹ năng2", "kỹ năng3", ...],
      "experience_years": "Số năm kinh nghiệm ước tính",
      "experience_summary": "Tóm tắt ngắn gọn về kinh nghiệm làm việc",
      "education": "Nền tảng giáo dục",
      "strengths": ["điểm mạnh1", "điểm mạnh2", ...],
      "weaknesses": ["điểm yếu1", "điểm yếu2", ...],
      "learning_path": {{
        "immediate": ["kỹ năng hoặc lĩnh vực cần học ngay lập tức"],
        "short_term": ["kỹ năng cho 3-6 tháng tới"],
        "long_term": ["kỹ năng cho 6-12 tháng"]
      }},
      "recommended_tasks": ["nhiệm vụ 1", "nhiệm vụ 2", ...]
    }}
 
    LƯU Ý: Tất cả giá trị trong JSON (strengths, weaknesses, learning_path, recommended_tasks) đều PHẢI là tiếng Việt.
    Chỉ trả về đối tượng JSON, không có văn bản bổ sung.
    """
    
    try:
        response = await GEMINI_MODEL.generate_content_async(prompt_template)
        raw_text = response.text.strip()
 
        match = re.search(r'```json\s*({.*?})\s*```|({.*?})', raw_text, re.DOTALL)
        if match:
            json_str = match.group(1) or match.group(2)
            analysis_data = json.loads(json_str)
            return JSONResponse(content=analysis_data)
        else:
            return JSONResponse(
                status_code=500,
                content={"error": "Could not parse AI response", "raw": raw_text}
            )
    except Exception as e:
        print(f"Error in CV analysis: {e}")
        return JSONResponse(status_code=500, content={"error": str(e)})

@router.post("/generate-cv")
async def generate_cv(data: CVGenerationRequest):
    """
    Generate a sample CV in markdown format based on user profile
    """
    skills_list = ", ".join(data.skills)
    achievements_list = "\n".join([f"- {a}" for a in data.achievements]) if data.achievements else "- [Add your achievements]"
    
    prompt_template = f"""
    You are an expert CV/resume writer.
    
    Create a professional CV in Markdown format for a candidate with the following profile using English:
    
    ROLE: {data.role}
    SKILLS: {skills_list}
    EXPERIENCE: {data.experience}
    EDUCATION: {data.education}
    ACHIEVEMENTS:
    {achievements_list}
    
    Generate a complete, professional CV in Markdown format with the following sections:
    - Header with name and contact (use placeholders)
    - Professional Summary
    - Skills
    - Work Experience
    - Education
    - Achievements
    - Additional relevant sections
    
    Make it ATS-friendly and professional. Use proper Markdown formatting.
    Return ONLY the Markdown content, no JSON, no code blocks.
    """
    
    if GEMINI_MODEL is None:
        return JSONResponse(
            status_code=503,
            content={
                "error": "Gemini is not configured. Set env var GOOGLE_API_KEY (or GEMINI_API_KEY) on the server.",
            },
        )

    try:
        response = await GEMINI_MODEL.generate_content_async(prompt_template)
        cv_markdown = response.text.strip()
        
        # Remove markdown code blocks if present
        cv_markdown = re.sub(r'^```markdown\s*', '', cv_markdown)
        cv_markdown = re.sub(r'```\s*$', '', cv_markdown)
        
        return JSONResponse(content={"cv_markdown": cv_markdown})
    except Exception as e:
        print(f"Error in CV generation: {e}")
        return JSONResponse(status_code=500, content={"error": str(e)})

@router.post("/generate-cv-docx")
async def generate_cv_docx(data: CVGenerationRequest):
    """
    Generate a CV in DOCX format based on user profile
    """
    skills_list = ", ".join(data.skills)
    achievements_list = "\n".join([f"- {a}" for a in data.achievements]) if data.achievements else "- [Add your achievements]"
    
    prompt_template = f"""
    You are an expert CV/resume writer.
    
    Create a professional CV for a candidate with the following profile using English:
    
    ROLE: {data.role}
    SKILLS: {skills_list}
    EXPERIENCE: {data.experience}
    EDUCATION: {data.education}
    ACHIEVEMENTS:
    {achievements_list}
    
    Generate a complete, professional CV with the following sections:
    - Header with [Your Full Name] and contact placeholders
    - Professional Summary (2-3 sentences)
    - Skills (list format)
    - Work Experience (with job titles, companies, dates, responsibilities)
    - Education
    - Achievements
    
    Make it ATS-friendly and professional. Use clear section headers.
    Return plain text content, no markdown syntax, no code blocks.
    """
    
    if GEMINI_MODEL is None:
        return JSONResponse(
            status_code=503,
            content={
                "error": "Gemini is not configured. Set env var GOOGLE_API_KEY (or GEMINI_API_KEY) on the server.",
            },
        )

    try:
        response = await GEMINI_MODEL.generate_content_async(prompt_template)
        cv_text = response.text.strip()
        
        # Create DOCX document
        doc = Document()
        
        # Parse the CV text and add to document
        lines = cv_text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Check if it's a section header (all caps or ends with colon)
            if line.isupper() or line.endswith(':'):
                # Add section header
                heading = doc.add_heading(line.rstrip(':'), level=2)
                heading.runs[0].font.color.rgb = RGBColor(0, 0, 139)
            else:
                # Add regular paragraph
                para = doc.add_paragraph(line)
                para.style = 'Normal'
        
        # Save to BytesIO
        docx_io = io.BytesIO()
        doc.save(docx_io)
        docx_io.seek(0)
        
        return StreamingResponse(
            docx_io,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=CV_Generated.docx"}
        )
    except Exception as e:
        print(f"Error in CV DOCX generation: {e}")
        return JSONResponse(status_code=500, content={"error": str(e)})

@router.post("/generate-questions")
async def generate_questions(data: QuestionGenerationRequest):
    skills_text = ", ".join(data.skills) if data.skills else "kỹ năng chuyên môn chung"
    role_text = data.role or data.field
    
    prompt_template = f"""Bạn là chuyên gia tuyển dụng nhân sự cấp cao. Hãy tạo các câu hỏi phỏng vấn cho hồ sơ sau:
 
TARGET ROLE: {role_text}
FIELD: {data.field}
KEY SKILLS: {skills_text}
 
CHỈ trả về một mảng JSON hợp lệ gồm 15-20 chuỗi (không có markdown, không có văn bản bên ngoài dấu ngoặc).
Mỗi câu hỏi PHẢI bắt đầu chính xác với một thẻ: [Background], [Situation], hoặc [Technical].
 
Định dạng ví dụ:
[
  "[Bối cảnh] Hãy kể cho tôi nghe về kinh nghiệm của bạn với việc phân tích dữ liệu.",
  "[Tình huống] Mô tả cách bạn xử lý một hạn chót khó khăn.",
  "[Kỹ thuật] Giải thích các khái niệm chính của học máy."
]
 
Đặt câu hỏi cụ thể cho vai trò và kỹ năng. CHỈ trả về mảng JSON, không trả về bất kỳ dữ liệu nào khác."""
    
    if GEMINI_MODEL is None:
        return JSONResponse(
            status_code=503,
            content={
                "error": "Gemini is not configured. Set env var GOOGLE_API_KEY (or GEMINI_API_KEY) on the server.",
            },
        )

    try:
        response = await GEMINI_MODEL.generate_content_async(prompt_template)
        raw_text = response.text.strip()
 
        # Remove markdown code blocks
        raw_text = re.sub(r'```json\s*', '', raw_text)
        raw_text = re.sub(r'```\s*', '', raw_text)
 
        # Find JSON array
        match = re.search(r'(\[.*\])', raw_text, re.DOTALL)
        if match:
            json_str = match.group(1)
            # Clean up common JSON issues
            json_str = re.sub(r',\s*([\]\}])', r'\1', json_str)  # Remove trailing commas
            json_str = json_str.replace("\n", " ").replace("\r", " ")  # Remove newlines
 
            questions = json.loads(json_str)
 
            if not isinstance(questions, list):
                raise ValueError("Phản hồi không phải là một danh sách")
 
            return JSONResponse(content={"questions": questions})
        else:
            return JSONResponse(
                status_code=500,
                content={"error": "Không thể tìm thấy mảng JSON từ AI.", "raw": raw_text[:500]}
            )
    except json.JSONDecodeError as e:
        return JSONResponse(
            status_code=500,
            content={"error": f"Lỗi phân tích JSON: {str(e)}", "raw": raw_text[:500]}
        )
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"error": f"Lỗi khi tạo câu hỏi: {str(e)}"}
        )